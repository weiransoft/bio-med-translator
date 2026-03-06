#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
通用 Word 文档翻译脚本
使用 AI 进行翻译，保持文档格式
"""

import sys
import re
from docx import Document


def translate_text(text, translations=None):
    """
    翻译文本
    
    Args:
        text: 要翻译的文本
        translations: 可选的翻译映射表
    
    Returns:
        翻译后的文本
    """
    if not text or not text.strip():
        return text
    
    # 如果提供了翻译映射表，先尝试匹配
    if translations and text.strip() in translations:
        return translations[text.strip()]
    
    # 处理带编号的章节标题 (如 "1. 目的 Purpose")
    pattern1 = r'^(\d+)\.\s*([\u4e00-\u9fa5]+)\s+([A-Za-z\s]+)$'
    match1 = re.match(pattern1, text.strip())
    if match1:
        number = match1.group(1)
        english = match1.group(3)
        return f"{number}. {english}"
    
    # 处理带编号的小节标题 (如 "5.1 总结 Summary")
    pattern2 = r'^(\d+\.\d+)\s*([\u4e00-\u9fa5]+)\s+([A-Za-z\s]+)$'
    match2 = re.match(pattern2, text.strip())
    if match2:
        number = match2.group(1)
        english = match2.group(3)
        return f"{number} {english}"
    
    # 如果提供了翻译映射表，进行部分匹配替换
    if translations:
        translated = text
        # 按长度降序排序，优先匹配长的
        sorted_translations = sorted(translations.items(), key=lambda x: len(x[0]), reverse=True)
        for key, value in sorted_translations:
            if key in translated:
                translated = translated.replace(key, value)
        return translated
    
    return text


def translate_document(input_path, output_path, translations=None):
    """
    翻译 Word 文档
    
    Args:
        input_path: 输入文件路径
        output_path: 输出文件路径
        translations: 可选的翻译映射表
    """
    print(f"正在读取文档: {input_path}")
    doc = Document(input_path)
    
    # 统计信息
    total_paragraphs = len(doc.paragraphs)
    total_tables = len(doc.tables)
    translated_paragraphs = 0
    translated_tables = 0
    
    print(f"文档包含 {total_paragraphs} 个段落和 {total_tables} 个表格")
    
    # 翻译段落
    print("\n正在翻译段落...")
    for i, para in enumerate(doc.paragraphs, 1):
        if para.text.strip():
            original_text = para.text
            translated_text = translate_text(original_text, translations)
            
            if translated_text != original_text:
                para.text = translated_text
                translated_paragraphs += 1
                if i <= 15 or i > total_paragraphs - 5:
                    print(f"  段落 {i}: {original_text[:50]}... -> {translated_text[:50]}...")
    
    print(f"段落翻译完成: {translated_paragraphs}/{total_paragraphs}")
    
    # 翻译表格
    print("\n正在翻译表格...")
    for table_idx, table in enumerate(doc.tables, 1):
        table_translated = False
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                if cell.text.strip():
                    original_text = cell.text
                    translated_text = translate_text(original_text, translations)
                    
                    if translated_text != original_text:
                        # 保留段落格式
                        for para in cell.paragraphs:
                            para.text = translate_text(para.text, translations)
                        table_translated = True
        
        if table_translated:
            translated_tables += 1
            print(f"  表格 {table_idx}: 已翻译")
    
    print(f"表格翻译完成: {translated_tables}/{total_tables}")
    
    # 翻译页眉和页脚
    print("\n正在翻译页眉和页脚...")
    header_footer_translated = 0
    for section in doc.sections:
        # 翻译页眉
        for para in section.header.paragraphs:
            if para.text.strip():
                original_text = para.text
                translated_text = translate_text(original_text, translations)
                if translated_text != original_text:
                    para.text = translated_text
                    header_footer_translated += 1
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        original_text = cell.text
                        translated_text = translate_text(original_text, translations)
                        if translated_text != original_text:
                            for para in cell.paragraphs:
                                para.text = translate_text(para.text, translations)
                            header_footer_translated += 1
        # 翻译页脚
        for para in section.footer.paragraphs:
            if para.text.strip():
                original_text = para.text
                translated_text = translate_text(original_text, translations)
                if translated_text != original_text:
                    para.text = translated_text
                    header_footer_translated += 1
        for table in section.footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        original_text = cell.text
                        translated_text = translate_text(original_text, translations)
                        if translated_text != original_text:
                            for para in cell.paragraphs:
                                para.text = translate_text(para.text, translations)
                            header_footer_translated += 1
    print(f"页眉页脚翻译完成: {header_footer_translated} 处")
    
    # 保存文档
    print(f"\n正在保存文档: {output_path}")
    doc.save(output_path)
    
    print(f"\n✅ 翻译完成！")
    print(f"  输入文件: {input_path}")
    print(f"  输出文件: {output_path}")
    print(f"  翻译段落: {translated_paragraphs}/{total_paragraphs}")
    print(f"  翻译表格: {translated_tables}/{total_tables}")
    print(f"  翻译页眉页脚: {header_footer_translated} 处")


if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("用法：python translate_docx.py <input.docx> <output.docx> [translations.json]")
        print("\n示例：")
        print("  python translate_docx.py input.docx output.docx")
        print("  python translate_docx.py input.docx output.docx translations.json")
        sys.exit(1)
    
    input_path = sys.argv[1]
    output_path = sys.argv[2]
    
    # 可选：加载翻译映射表
    translations = None
    if len(sys.argv) > 3:
        import json
        with open(sys.argv[3], 'r', encoding='utf-8') as f:
            translations = json.load(f)
        print(f"已加载翻译映射表: {sys.argv[3]}")
    
    translate_document(input_path, output_path, translations)
