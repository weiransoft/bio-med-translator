#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
验证 Word 文档翻译结果
检查是否还有中文残留
"""

import sys
import re
from docx import Document


def check_chinese(text):
    """检查文本中是否包含中文字符"""
    pattern = re.compile(r'[\u4e00-\u9fa5]')
    return pattern.search(text) is not None


def verify_document(docx_path):
    """
    验证文档翻译是否完整
    
    Args:
        docx_path: Word 文档路径
    
    Returns:
        (is_complete, chinese_items) 元组
    """
    print(f"正在验证文档: {docx_path}")
    doc = Document(docx_path)
    
    chinese_items = []
    
    # 检查段落
    for i, para in enumerate(doc.paragraphs, 1):
        if para.text.strip() and check_chinese(para.text):
            chinese_items.append({
                'type': '段落',
                'index': i,
                'text': para.text[:100] + ('...' if len(para.text) > 100 else '')
            })
    
    # 检查表格
    for table_idx, table in enumerate(doc.tables, 1):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                if cell.text.strip() and check_chinese(cell.text):
                    chinese_items.append({
                        'type': f'表格 {table_idx}',
                        'index': f'行{row_idx+1},列{cell_idx+1}',
                        'text': cell.text[:100] + ('...' if len(cell.text) > 100 else '')
                    })
    
    # 检查页眉页脚
    for section_idx, section in enumerate(doc.sections, 1):
        # 页眉
        for para_idx, para in enumerate(section.header.paragraphs, 1):
            if para.text.strip() and check_chinese(para.text):
                chinese_items.append({
                    'type': f'页眉 (Section {section_idx})',
                    'index': f'段落 {para_idx}',
                    'text': para.text[:100] + ('...' if len(para.text) > 100 else '')
                })
        # 页脚
        for para_idx, para in enumerate(section.footer.paragraphs, 1):
            if para.text.strip() and check_chinese(para.text):
                chinese_items.append({
                    'type': f'页脚 (Section {section_idx})',
                    'index': f'段落 {para_idx}',
                    'text': para.text[:100] + ('...' if len(para.text) > 100 else '')
                })
    
    is_complete = len(chinese_items) == 0
    
    # 输出结果
    print(f"\n{'='*60}")
    if is_complete:
        print("✅ 验证通过！文档已完全翻译为英文。")
    else:
        print(f"❌ 发现 {len(chinese_items)} 处中文内容未翻译：")
        print(f"{'='*60}")
        for item in chinese_items[:20]:  # 最多显示20条
            print(f"\n【{item['type']}】位置: {item['index']}")
            print(f"  内容: {item['text']}")
        
        if len(chinese_items) > 20:
            print(f"\n... 还有 {len(chinese_items) - 20} 处未显示 ...")
    print(f"{'='*60}")
    
    return is_complete, chinese_items


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("用法：python verify_translation.py <docx_file>")
        print("\n示例：")
        print("  python verify_translation.py translated.docx")
        sys.exit(1)
    
    docx_path = sys.argv[1]
    is_complete, chinese_items = verify_document(docx_path)
    
    # 返回退出码
    sys.exit(0 if is_complete else 1)
